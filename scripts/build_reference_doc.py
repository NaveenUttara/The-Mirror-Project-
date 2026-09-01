from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\Uttara ERP\mirror-project\docs\The_Mirror_Project_Technical_Reference.docx")

GREEN = "087F5B"
DARK_GREEN = "05523E"
INK = "12212B"
MUTED = "667781"
LIGHT_GREEN = "EAF5F1"
LIGHT_GRAY = "F2F4F7"
BORDER = "CFDDD8"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_in):
    total_dxa = 9360
    widths_dxa = [round(value * 1440) for value in widths_in]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(widths_in[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_table(table, header=True, font_size=8.7):
    if header:
        set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0 and header:
                set_cell_shading(cell, GREEN)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F7FAF9")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=font_size,
                        color=WHITE if row_index == 0 and header else INK,
                        bold=True if row_index == 0 and header else None,
                    )


def add_table(doc, headers, rows, widths, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = str(value)
    set_table_geometry(table, widths)
    format_table(table, font_size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        text = text[len(bold_lead):]
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9, color=INK)
    return p


def add_note(doc, label, text):
    p = doc.add_paragraph(style="Note Box")
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, bold=True, color=DARK_GREEN)
    run = p.add_run(text)
    set_run_font(run, color=INK)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, 18, 10),
        2: (13, 14, 7),
        3: (12, 10, 5),
    }
    for level, (size, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(GREEN if level < 3 else DARK_GREEN)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.05
    p_pr = code._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)

    note = styles.add_style("Note Box", 1)
    note.font.name = "Calibri"
    note.font.size = Pt(10.5)
    note.paragraph_format.left_indent = Inches(0.15)
    note.paragraph_format.right_indent = Inches(0.15)
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(10)
    note.paragraph_format.line_spacing = 1.15
    note_p_pr = note._element.get_or_add_pPr()
    note_shd = OxmlElement("w:shd")
    note_shd.set(qn("w:fill"), LIGHT_GREEN)
    note_p_pr.append(note_shd)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("THE MIRROR PROJECT   |   TECHNICAL REFERENCE")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(110)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("THE MIRROR PROJECT")
    set_run_font(r, size=13, color=GREEN, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Backend & Database\nReference Guide")
    set_run_font(r, size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FRSCMP Oracle Integration and Visual Studio Code File Map")
    set_run_font(r, size=14, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(46)
    r = p.add_run("See it. Report it. Track it. Fix it.")
    set_run_font(r, size=12, color=GREEN, bold=True)

    add_table(
        doc,
        ["Document", "Details"],
        [
            ("Purpose", "Explain what each Oracle table and VS Code file does"),
            ("Audience", "Developers, database administrators, and project reviewers"),
            ("Environment", "Next.js 16, Node.js, Oracle FRSCMP, IIS reverse proxy"),
            ("Prepared", "1 September 2026"),
        ],
        [1.45, 5.05],
        font_size=9.2,
    )
    add_note(
        doc,
        "Confidentiality",
        "This guide never includes database passwords, JWT secrets, or other private environment values.",
    )
    doc.add_page_break()


def add_overview(doc):
    add_heading(doc, "1. System Overview", 1)
    add_body(
        doc,
        "The Mirror Project is a Karnataka pothole reporting and accountability application. Citizens authenticate using OTP, report potholes with evidence and location, and track each case until repair and citizen confirmation.",
    )
    add_heading(doc, "1.1 Architecture", 2)
    add_code(
        doc,
        "Citizen browser\n"
        "    -> Next.js frontend (app/page.tsx)\n"
        "    -> Next.js API routes (app/api)\n"
        "    -> Oracle connection pool (lib/db.ts)\n"
        "    -> FRSCMP Oracle database (MIRROR_* objects)\n"
        "    -> Object/file storage for photographs\n"
        "    -> SMS provider for OTP delivery",
    )
    add_note(
        doc,
        "Deployment rule",
        "Local development uses http://localhost:3000. The root index.html is a legacy prototype and is not served by Next.js.",
    )

    add_heading(doc, "1.2 Current Implementation Status", 2)
    add_table(
        doc,
        ["Area", "Status", "Meaning"],
        [
            ("Next.js application", "Working", "Homepage and API routes compile and run"),
            ("Oracle authentication", "Verified", "The application can authenticate to FRSCMP using Thick mode"),
            ("Oracle migration", "Prepared", "The MIRROR_* table script exists but must be run by the DBA"),
            ("OTP backend", "Partial", "Request/verification code exists; rate limiting and sessions remain"),
            ("Pothole reporting", "Planned", "Database structures exist; form and API are not implemented"),
            ("IIS publishing", "Prepared", "Reverse-proxy template exists; production deployment remains"),
        ],
        [1.55, 1.15, 3.8],
    )
    add_heading(doc, "1.3 Database Isolation Principle", 2)
    add_body(
        doc,
        "All new database objects use a MIRROR_ prefix. This keeps the project separate from legacy FRSCMP tables while using the same approved Oracle database environment.",
    )
    doc.add_page_break()


TABLES = [
    (
        "MIRROR_USERS",
        "Stores citizens, officers, service providers, administrators, and super administrators.",
        [
            ("id", "NUMBER(19)", "Yes", "Primary key; generated by MIRROR_USERS_SEQ"),
            ("name", "NVARCHAR2(150)", "Yes", "Full display name"),
            ("phone", "VARCHAR2(20)", "Yes", "Unique verified login identity"),
            ("email", "VARCHAR2(254)", "No", "Notification/contact email"),
            ("role", "VARCHAR2(30)", "Yes", "citizen, officer, service_provider, admin, or super_admin"),
            ("phone_verified", "NUMBER(1)", "Yes", "0 = no; 1 = yes"),
            ("created_at", "TIMESTAMP", "Yes", "Creation time"),
            ("updated_at", "TIMESTAMP", "Yes", "Last update time"),
        ],
    ),
    (
        "MIRROR_OTP_REQUESTS",
        "Stores short-lived hashed OTP challenges. Plain OTP values must never be stored.",
        [
            ("id", "NUMBER(19)", "Yes", "Primary key; generated by MIRROR_OTP_REQ_SEQ"),
            ("phone", "VARCHAR2(20)", "Yes", "Phone number receiving the OTP"),
            ("otp_hash", "VARCHAR2(255)", "Yes", "Salted OTP hash"),
            ("expires_at", "TIMESTAMP", "Yes", "Expiry time"),
            ("attempts", "NUMBER(3)", "Yes", "Failed verification count"),
            ("consumed_at", "TIMESTAMP", "No", "Successful-use time"),
            ("request_ip_hash", "VARCHAR2(64)", "No", "Privacy-preserving requester identifier"),
            ("created_at", "TIMESTAMP", "Yes", "Request creation time"),
        ],
    ),
    (
        "MIRROR_SESSIONS",
        "Stores revocable server-side sessions after successful authentication.",
        [
            ("id", "NUMBER(19)", "Yes", "Primary key; generated by MIRROR_SESSIONS_SEQ"),
            ("user_id", "NUMBER(19)", "Yes", "Foreign key to MIRROR_USERS"),
            ("token_hash", "VARCHAR2(64)", "Yes", "Unique hash of the session token"),
            ("expires_at", "TIMESTAMP", "Yes", "Session expiry"),
            ("revoked_at", "TIMESTAMP", "No", "Logout/revocation time"),
            ("created_at", "TIMESTAMP", "Yes", "Session creation time"),
        ],
    ),
    (
        "MIRROR_POTHOLES",
        "Represents one physical pothole. Multiple citizen reports may refer to the same pothole.",
        [
            ("id", "NUMBER(19)", "Yes", "Primary key; generated by MIRROR_POTHOLES_SEQ"),
            ("public_id", "VARCHAR2(30)", "Yes", "Unique public pothole identifier"),
            ("latitude", "NUMBER(10,7)", "Yes", "GPS latitude from -90 to 90"),
            ("longitude", "NUMBER(10,7)", "Yes", "GPS longitude from -180 to 180"),
            ("severity", "VARCHAR2(20)", "Yes", "low, medium, high, or critical"),
            ("current_status", "VARCHAR2(30)", "Yes", "Current workflow status"),
            ("created_at", "TIMESTAMP", "Yes", "Creation time"),
            ("updated_at", "TIMESTAMP", "Yes", "Last status/data change"),
        ],
    ),
    (
        "MIRROR_REPORTS",
        "Stores each citizen submission and links it to a physical pothole and reporting citizen.",
        [
            ("id", "NUMBER(19)", "Yes", "Primary key; generated by MIRROR_REPORTS_SEQ"),
            ("report_id", "VARCHAR2(30)", "Yes", "Unique public ID such as MIR-POT-000001"),
            ("pothole_id", "NUMBER(19)", "Yes", "Foreign key to MIRROR_POTHOLES"),
            ("citizen_id", "NUMBER(19)", "Yes", "Foreign key to MIRROR_USERS"),
            ("description", "CLOB", "No", "Citizen description/context"),
            ("submitted_at", "TIMESTAMP", "Yes", "Submission time"),
        ],
    ),
    (
        "MIRROR_REPORT_PHOTOS",
        "Stores photograph metadata and private storage references; image binary data is stored outside Oracle.",
        [
            ("id", "NUMBER(19)", "Yes", "Primary key; generated by MIRROR_PHOTOS_SEQ"),
            ("report_id", "NUMBER(19)", "Yes", "Foreign key to MIRROR_REPORTS"),
            ("object_key", "VARCHAR2(500)", "Yes", "Unique private storage key/path"),
            ("mime_type", "VARCHAR2(100)", "Yes", "Validated image MIME type"),
            ("file_size", "NUMBER(19)", "Yes", "File size in bytes"),
            ("evidence_type", "VARCHAR2(20)", "Yes", "before or after"),
            ("created_at", "TIMESTAMP", "Yes", "Metadata creation time"),
        ],
    ),
    (
        "MIRROR_STATUS_HISTORY",
        "Provides the immutable accountability trail for every report status change.",
        [
            ("id", "NUMBER(19)", "Yes", "Primary key; generated by MIRROR_STATUS_SEQ"),
            ("report_id", "NUMBER(19)", "Yes", "Foreign key to MIRROR_REPORTS"),
            ("old_status", "VARCHAR2(30)", "No", "Previous status"),
            ("new_status", "VARCHAR2(30)", "Yes", "New status"),
            ("changed_by", "NUMBER(19)", "No", "Foreign key to MIRROR_USERS"),
            ("comments", "VARCHAR2(1000)", "No", "Reason/evidence note"),
            ("created_at", "TIMESTAMP", "Yes", "Change time"),
        ],
    ),
    (
        "MIRROR_AUDIT_LOGS",
        "Records security-sensitive and administrative activity for traceability.",
        [
            ("id", "NUMBER(19)", "Yes", "Primary key; generated by MIRROR_AUDIT_SEQ"),
            ("actor_user_id", "NUMBER(19)", "No", "Foreign key to MIRROR_USERS"),
            ("action", "VARCHAR2(100)", "Yes", "Action performed"),
            ("entity_type", "VARCHAR2(50)", "Yes", "Affected object type"),
            ("entity_id", "VARCHAR2(100)", "No", "Affected object identifier"),
            ("details", "CLOB", "No", "Additional audit information"),
            ("created_at", "TIMESTAMP", "Yes", "Audit event time"),
        ],
    ),
]


def add_database_reference(doc):
    add_heading(doc, "2. Oracle Database Reference", 1)
    add_body(
        doc,
        "The migration defines eight tables, eight sequences, supporting indexes, validation constraints, and foreign-key relationships. The migration is Oracle 11g-compatible and must be executed only once using Run Script (F5) in Oracle SQL Developer.",
    )
    add_heading(doc, "2.1 Object Summary", 2)
    add_table(
        doc,
        ["Table", "Primary purpose", "Implementation status"],
        [(name, purpose, "Defined in migration") for name, purpose, _ in TABLES],
        [2.05, 3.55, 0.9],
        font_size=8.2,
    )
    add_heading(doc, "2.2 Relationships", 2)
    add_code(
        doc,
        "MIRROR_USERS\n"
        "  +-- MIRROR_SESSIONS\n"
        "  +-- MIRROR_REPORTS\n"
        "  +-- MIRROR_STATUS_HISTORY\n"
        "  +-- MIRROR_AUDIT_LOGS\n\n"
        "MIRROR_POTHOLES\n"
        "  +-- MIRROR_REPORTS\n"
        "        +-- MIRROR_REPORT_PHOTOS\n"
        "        +-- MIRROR_STATUS_HISTORY",
    )
    add_note(
        doc,
        "Important",
        "MIRROR_STATUS_HISTORY should be append-only during normal operation. A status change creates a new history row instead of erasing the previous state.",
    )
    doc.add_page_break()

    for index, (name, purpose, columns) in enumerate(TABLES, start=1):
        add_heading(doc, f"2.{index + 2} {name}", 2)
        add_body(doc, purpose)
        add_table(
            doc,
            ["Column", "Oracle type", "Required", "Purpose / rule"],
            columns,
            [1.35, 1.35, 0.75, 3.05],
            font_size=8.35,
        )
        if index in (2, 4, 6):
            doc.add_page_break()

    add_heading(doc, "2.11 Sequences and Indexes", 2)
    add_table(
        doc,
        ["Object", "Used for"],
        [
            ("MIRROR_USERS_SEQ", "MIRROR_USERS.id"),
            ("MIRROR_OTP_REQ_SEQ", "MIRROR_OTP_REQUESTS.id"),
            ("MIRROR_SESSIONS_SEQ", "MIRROR_SESSIONS.id"),
            ("MIRROR_POTHOLES_SEQ", "MIRROR_POTHOLES.id"),
            ("MIRROR_REPORTS_SEQ", "MIRROR_REPORTS.id"),
            ("MIRROR_PHOTOS_SEQ", "MIRROR_REPORT_PHOTOS.id"),
            ("MIRROR_STATUS_SEQ", "MIRROR_STATUS_HISTORY.id"),
            ("MIRROR_AUDIT_SEQ", "MIRROR_AUDIT_LOGS.id"),
            ("IX_MIRROR_OTP_PHONE_EXP", "OTP lookup by phone and expiry"),
            ("IX_MIRROR_SESSION_USER", "Active session lookup by user"),
            ("IX_MIRROR_REPORT_USER", "Citizen report listing"),
            ("IX_MIRROR_STATUS_REPORT", "Report history timeline"),
        ],
        [2.55, 3.95],
    )
def add_code_reference(doc):
    add_heading(doc, "3. Visual Studio Code File Map", 1)
    add_body(
        doc,
        "The application uses the Next.js App Router. The files below control the visible interface, authentication APIs, Oracle connectivity, validation, build behavior, and IIS proxying.",
    )
    add_table(
        doc,
        ["File", "What it does", "When to change it"],
        [
            ("app/page.tsx", "Main client interface; navigation, phone/OTP state, API calls, dashboard prototype", "UI, text, colours, buttons, frontend flows"),
            ("app/layout.tsx", "Root HTML layout, metadata, fonts, global wrapper", "Page title, description, shared layout"),
            ("app/globals.css", "Global CSS and Tailwind import", "Global typography/background rules"),
            ("app/api/auth/request-otp/route.ts", "Validates phone, creates hashed OTP, inserts MIRROR_OTP_REQUESTS, optionally calls SMS", "OTP generation, expiry, rate limiting, SMS"),
            ("app/api/auth/verify-otp/route.ts", "Loads latest OTP, verifies hash/expiry, creates/loads user, commits transaction, returns login token", "OTP attempts, registration, sessions, cookies"),
            ("lib/db.ts", "Loads Oracle Client Thick mode and maintains the Oracle connection pool", "Oracle pool options and environment validation"),
            ("next.config.ts", "Keeps native oracledb outside Turbopack's virtual filesystem", "Next.js server/build configuration"),
            (".env.local", "Private development configuration", "Oracle credentials, connection string, JWT secret, Instant Client directory"),
            ("package.json", "Scripts and dependency list", "Commands or npm dependencies"),
            ("database/oracle/001_mirror_foundation.sql", "Creates all MIRROR_* tables, sequences, indexes, and constraints", "Database design changes approved by DBA"),
            ("scripts/check-oracle.mjs", "Read-only Oracle authentication and required-table checker", "Database diagnostics"),
            ("web.config", "IIS reverse proxy template to internal Next.js port 3000", "Production IIS deployment"),
            ("index.html", "Legacy standalone prototype; not used by Next.js", "Do not use for active application changes"),
        ],
        [2.0, 2.8, 1.7],
        font_size=8.1,
    )

    add_heading(doc, "3.1 Confidential Environment Variables", 2)
    add_table(
        doc,
        ["Variable", "Purpose", "Browser-visible?"],
        [
            ("ORACLE_USER", "Oracle login user", "No"),
            ("ORACLE_PASSWORD", "Oracle login password", "No"),
            ("ORACLE_CONNECT_STRING", "Oracle host, listener port, and service", "No"),
            ("ORACLE_CLIENT_LIB_DIR", "Folder containing the supported 64-bit oci.dll", "No"),
            ("JWT_SECRET", "Signs authentication tokens during the current implementation", "No"),
            ("SMS_API_KEY", "Optional SMS provider credential", "No"),
        ],
        [2.15, 3.35, 1.0],
    )
    add_note(
        doc,
        "Security",
        "Never commit .env.local or prefix these values with NEXT_PUBLIC_. Any NEXT_PUBLIC_ value may be included in browser JavaScript.",
    )
def add_workflows(doc):
    add_heading(doc, "4. Application Action Map", 1)
    add_heading(doc, "4.1 Current OTP Workflow", 2)
    add_table(
        doc,
        ["User action", "Code path", "Database action", "Current status"],
        [
            ("Open website", "app/page.tsx", "None", "Working"),
            ("Enter phone and Send OTP", "request-otp/route.ts", "Insert MIRROR_OTP_REQUESTS", "Implemented; table must exist"),
            ("Receive SMS", "SMS provider call", "OTP row retained", "Provider configuration pending"),
            ("Enter OTP and Verify", "verify-otp/route.ts", "Read/delete OTP; read/create MIRROR_USERS", "Implemented"),
            ("Open dashboard", "app/page.tsx", "None", "Prototype only"),
        ],
        [1.35, 1.65, 2.2, 1.3],
        font_size=8.2,
    )

    add_heading(doc, "4.2 Planned Pothole Reporting Workflow", 2)
    for text in (
        "Citizen signs in through OTP and receives a secure session.",
        "Citizen captures a photograph, GPS coordinates, severity, and description.",
        "Backend validates the session, coordinates, image metadata, and payload.",
        "Backend detects a nearby existing pothole or creates MIRROR_POTHOLES.",
        "Backend creates MIRROR_REPORTS with a public report ID.",
        "Photo metadata is stored in MIRROR_REPORT_PHOTOS; image bytes remain in private storage.",
        "Initial Reported state is appended to MIRROR_STATUS_HISTORY.",
        "Officer, service-provider, repair-evidence, citizen-confirmation, and escalation stages follow.",
    ):
        add_number(doc, text)

    add_heading(doc, "4.3 Database Migration Workflow", 2)
    add_code(
        doc,
        "1. Back up and obtain DBA approval\n"
        "2. Connect to the approved FRSCMP schema in Oracle SQL Developer\n"
        "3. Open database/oracle/001_mirror_foundation.sql\n"
        "4. Execute once with F5 (Run Script)\n"
        "5. Run npm run db:check\n"
        "6. Confirm MIRROR_USERS and MIRROR_OTP_REQUESTS are PRESENT",
    )
    add_note(
        doc,
        "Oracle DDL",
        "Table and sequence creation automatically commits. Do not rerun the complete migration after objects have been created.",
    )
def add_operations(doc):
    add_heading(doc, "5. Developer Operations", 1)
    add_heading(doc, "5.1 Common Commands", 2)
    add_table(
        doc,
        ["Command", "Purpose"],
        [
            ("npm run dev", "Start development server with automatic refresh"),
            ("npm run db:check", "Verify Oracle authentication and required tables"),
            ("npm run lint", "Check code style and common defects"),
            ("npx tsc --noEmit", "Validate TypeScript without producing output files"),
            ("npm run build", "Create and validate the production build"),
            ("npm run start", "Run the previously built production application"),
        ],
        [2.2, 4.3],
    )
    add_heading(doc, "5.2 Local Development Rules", 2)
    add_bullet(doc, "Keep only one Next.js server on port 3000.")
    add_bullet(doc, "Use http://localhost:3000, not index.html or IIS Express.")
    add_bullet(doc, "Use npm run dev while changing features; use npm run start only for production-build testing.")
    add_bullet(doc, "Restart the server after changing .env.local or next.config.ts.")
    add_bullet(doc, "Run lint, TypeScript, and build checks after every completed feature.")

    add_heading(doc, "5.3 IIS Deployment Boundary", 2)
    add_body(
        doc,
        "The web.config file is a reverse-proxy template. IIS accepts external HTTP traffic and forwards it to a Next.js process listening on 127.0.0.1:3000. The existing FRSCMP ASP.NET login on port 8080 must not be overwritten without explicit administrator approval.",
    )
    add_code(
        doc,
        "External browser -> IIS binding -> web.config rewrite\n"
        "                 -> http://127.0.0.1:3000 -> Next.js",
    )

    add_heading(doc, "5.4 Security Requirements", 2)
    add_bullet(doc, "Rotate any credential or JWT secret that has been shared outside the protected environment.")
    add_bullet(doc, "Replace localStorage JWT storage with HttpOnly, Secure, SameSite cookies and MIRROR_SESSIONS.")
    add_bullet(doc, "Add OTP request limits per phone and IP, maximum verification attempts, and resend cooldowns.")
    add_bullet(doc, "Return generic API errors to users; retain detailed Oracle errors only in protected server logs.")
    add_bullet(doc, "Keep citizen photographs private and store only validated metadata/object keys in Oracle.")
    add_bullet(doc, "Use least-privilege Oracle permissions and preserve the MIRROR_ object prefix.")
def add_roadmap(doc):
    add_heading(doc, "6. Remaining Backend Roadmap", 1)
    add_table(
        doc,
        ["Priority", "Work item", "Definition of done"],
        [
            ("1", "Execute Oracle migration", "All MIRROR_* tables and sequences exist; db:check passes"),
            ("2", "Harden OTP", "Rate limits, attempts, expiry, one-time consumption, checked SMS response"),
            ("3", "Secure sessions", "HttpOnly cookie and MIRROR_SESSIONS replace localStorage JWT"),
            ("4", "Citizen registration", "First-time name/email flow and protected dashboard"),
            ("5", "Report submission", "Photo, GPS, severity, description, public report ID"),
            ("6", "Photo storage", "Private object storage, validation, metadata, before/after evidence"),
            ("7", "Tracking", "Citizen-owned report list and immutable status timeline"),
            ("8", "Officer workflow", "Jurisdiction, assignment, work orders, repair verification"),
            ("9", "Citizen confirmation", "Repaired/still damaged, reopen, escalation, closure"),
            ("10", "Production deployment", "Windows service, IIS reverse proxy, HTTPS, monitoring and backups"),
        ],
        [0.65, 1.75, 4.1],
        font_size=8.5,
    )
    add_heading(doc, "6.1 Acceptance Checklist", 2)
    for item in (
        "Oracle migration executed once with no errors.",
        "npm run db:check reports authentication success and required tables present.",
        "npm run lint, npx tsc --noEmit, and npm run build all pass.",
        "OTP succeeds without returning sensitive server errors.",
        "No confidential values are tracked by Git.",
        "Backup, HTTPS, least-privilege database access, and monitoring are documented before public launch.",
    ):
        add_bullet(doc, item)

    add_note(
        doc,
        "Source of truth",
        "For database structure, use database/oracle/001_mirror_foundation.sql. For application behavior, use the active app/, lib/, next.config.ts, and package.json files rather than the legacy index.html prototype.",
    )


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "The Mirror Project Backend and Database Reference Guide"
    doc.core_properties.subject = "FRSCMP Oracle integration and Visual Studio Code file map"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "Mirror Project, Oracle, FRSCMP, Next.js, pothole reporting"

    add_cover(doc)
    add_overview(doc)
    add_database_reference(doc)
    add_code_reference(doc)
    add_workflows(doc)
    add_operations(doc)
    add_roadmap(doc)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

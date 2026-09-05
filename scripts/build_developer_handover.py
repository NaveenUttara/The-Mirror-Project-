from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "The_Mirror_Project_Developer_Handover.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9D9D9", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, name="Arial", size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    for style_name, size in (("Heading 1", 18), ("Heading 2", 13), ("Heading 3", 11)):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)

    styles["List Bullet"].font.name = "Arial"
    styles["List Number"].font.name = "Arial"


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, name="Consolas", size=9, color="333333")
    return paragraph


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for index, heading in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, "1F4E78")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(heading)
        set_run_font(run, size=9.5, bold=True, color="FFFFFF")
        if widths:
            cell.width = Inches(widths[index])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for index, value in enumerate(values):
            cell = cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_shading(cell, "F2F6FA")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9)
            if widths:
                cell.width = Inches(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p_pr = title._p.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)
    title.add_run("The Mirror Project Complete Development and Handover Guide")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(8)
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("System flow services file responsibilities operating procedures and remaining work")
    set_run_font(run, size=15, color="404040")

    meta_rows = [
        ("Prepared", "5 September 2026"),
        ("Repository", "The Mirror Project"),
        ("Current phase", "Backend migration in progress"),
        ("Target backend", "MedusaJS on Railway"),
        ("Target data and storage services", "Neon PostgreSQL Upstash Redis and Cloudflare R2"),
        ("Primary domain", "themirror-project.in"),
    ]
    add_table(doc, ["Item", "Value"], meta_rows, [1.65, 5.05])

    doc.add_paragraph(
        "This document gives a new developer enough context to run the existing application, "
        "understand the full request flow, identify the service and file responsible for each action, "
        "protect credentials, and continue the work in the correct order. The existing Next.js "
        "application still uses Oracle for its working authentication and report APIs. The MedusaJS "
        "backend now connects to Neon PostgreSQL and Upstash Redis, but the frontend has not yet been "
        "switched to MedusaJS. The working Next.js application stores new private photographs in "
        "Cloudflare R2 and serves them through an authenticated application API."
    )
    paragraph = doc.add_paragraph()
    lead = paragraph.add_run("Main conclusion  ")
    set_run_font(lead, bold=True)
    paragraph.add_run(
        "The PostgreSQL schema and Redis connection are ready. The migration is not production complete "
        "until Medusa API routes, production photo-storage deployment, frontend integration, automated testing, and "
        "Railway deployment are finished."
    )

    add_page_break(doc)
    add_heading(doc, "Contents", 1)
    for item in (
        "1 Current system status",
        "2 Existing and target architecture",
        "3 Development tools services and purchase status",
        "4 Important files and responsibilities",
        "5 Changes added in the Medusa migration",
        "6 PostgreSQL data model",
        "7 Environment variables and secrets",
        "8 Running and verifying the project",
        "9 Remaining implementation work",
        "10 Deployment and release process",
        "11 Developer handover checklist",
        "12 Troubleshooting reference",
    ):
        doc.add_paragraph(item)

    add_heading(doc, "1 Current system status", 1)
    add_table(
        doc,
        ["Area", "Status", "What this means"],
        [
            ("Next.js user interface", "Working", "Runs from the repository root and is the current citizen interface."),
            ("Existing API routes", "Working with Oracle", "OTP verification and report operations still call the Oracle database."),
            ("MedusaJS backend", "Scaffolded", "A separate backend workspace exists under backend and compiles successfully."),
            ("Neon PostgreSQL", "Connected", "Standard Medusa migrations and eight Mirror tables were applied successfully."),
            ("Upstash Redis", "Connected", "The backend TLS connection was verified with a successful PING and PONG response."),
            ("Cloudflare R2", "Implemented; credentials pending verification", "New files use a private R2 bucket; authenticated retrieval includes a fallback for three earlier local files."),
            ("Medusa Mirror APIs", "Not implemented", "The new tables cannot yet receive citizen input through Medusa."),
            ("Frontend switch", "Not started", "The browser still calls local Next.js API routes rather than Railway Medusa endpoints."),
            ("Railway deployment", "Not started", "The Medusa service does not yet have a public backend URL."),
            ("Git state", "Action required", "The backend directory is untracked and must be committed only after secrets and tests are checked."),
        ],
        [1.5, 1.55, 3.65],
    )

    add_heading(doc, "Immediate next action", 2)
    doc.add_paragraph(
        "Implement the Medusa OTP session and report APIs against Neon and Upstash. Keep the working "
        "Oracle routes and R2 photograph storage available until the replacement APIs and their "
        "production storage design are tested."
    )

    add_heading(doc, "2 Existing and target architecture", 1)
    add_heading(doc, "Existing working path", 2)
    add_code(doc, "Browser -> Next.js on Vercel or localhost -> Next.js API routes -> Oracle database")
    doc.add_paragraph(
        "This path remains the rollback path. Files at the repository root contain the user interface, "
        "OTP routes, report route, JWT helper, and Oracle connector. Removing Oracle code now would break "
        "the working application before the Medusa replacement is ready."
    )

    add_heading(doc, "Target production path", 2)
    add_code(doc, "Browser -> Next.js on Vercel -> MedusaJS on Railway -> Neon PostgreSQL")
    add_code(doc, "                                      |-> Upstash Redis")
    add_code(doc, "                                      |-> Cloudflare R2 private photograph storage")
    add_table(
        doc,
        ["Service", "Responsibility", "Data stored"],
        [
            ("Vercel", "Hosts the existing Next.js frontend", "Built frontend assets and server-rendered pages"),
            ("Railway", "Runs the MedusaJS backend continuously", "Application runtime and logs"),
            ("Neon", "Stores structured PostgreSQL records", "Users OTP records sessions potholes reports photo metadata status history and audits"),
            ("Upstash", "Supplies Redis for Medusa infrastructure", "Cache events workflow state and distributed locks"),
            ("Cloudflare R2", "Stores private report photographs", "Image objects; access is proxied through authenticated application APIs"),
            ("BigRock", "Manages the purchased domain DNS", "DNS records only"),
        ],
        [1.25, 2.2, 3.25],
    )

    add_heading(doc, "3 Development tools services and purchase status", 1)
    add_table(
        doc,
        ["Tool or service", "Purpose", "Current commercial status"],
        [
            ("VS Code and PowerShell", "Edit files and run local commands", "Installed development tools"),
            ("Git and GitHub", "Version control and source repository", "Repository active"),
            ("Node.js and npm", "Run Next.js and MedusaJS", "Open source runtime and package manager"),
            ("Next.js React and TypeScript", "Citizen website and current API routes", "Open source dependencies"),
            ("MedusaJS 2.20.1", "Replacement backend platform", "Open source and no Medusa Cloud plan purchased"),
            ("Oracle FRSCMP", "Current users OTP reports and status data", "Existing company database and temporary rollback path"),
            ("TOAD and Oracle Instant Client 19.32", "Oracle administration and native database connectivity", "Existing local tools"),
            ("Neon PostgreSQL", "Target relational database", "Free project active"),
            ("Upstash Redis", "Cache events workflows and locks", "Free database active and connection verified"),
            ("Cloudflare R2", "Private report photograph objects", "Free allowance expected initially; bucket and API credentials required"),
            ("Railway", "Target MedusaJS hosting", "Pending service creation and paid usage"),
            ("Vercel", "Hosts the Next.js frontend and HTTPS domain", "Project active on the shown Hobby tier"),
            ("BigRock themirror-project.in", "Primary production domain and DNS", "Registered for three years according to the purchase account"),
            ("BigRock themirror-project.info", "Secondary or temporary domain", "Registered for one year and may be retired when no longer needed"),
        ],
        [1.7, 2.5, 2.5],
    )
    doc.add_paragraph(
        "A domain registration supplies the name and DNS control; it does not supply application hosting, "
        "a public server, or a database. Vercel hosts the frontend, Railway will host MedusaJS, Neon stores "
        "relational records, Upstash supports Medusa infrastructure, and Cloudflare R2 stores private "
        "photograph objects. The application API checks citizen ownership before returning an image."
    )

    add_heading(doc, "4 Important files and responsibilities", 1)
    add_table(
        doc,
        ["File or folder", "Action performed", "When to change it"],
        [
            ("app/page.tsx", "Controls home login OTP profile dashboard report photo GPS success and impact screens", "Change citizen interface behavior and browser API calls"),
            ("app/mirror.css", "Styles the Mirror user interface and responsive layout", "Change visual design spacing colors and mobile layout"),
            ("app/layout.tsx and app/globals.css", "Set page metadata and global application styles", "Change global page shell metadata or shared styles"),
            ("app/api/auth/request-otp/route.ts", "Validates the phone creates a hashed OTP record and currently exposes temporary OTP 123456 when SMS is absent", "Replace temporary OTP delivery when the SMS provider is ready"),
            ("app/api/auth/verify-otp/route.ts", "Verifies OTP creates or updates the user profile and returns a seven day JWT", "Change login profile validation or token behavior"),
            ("app/api/reports/route.ts", "Authenticates users validates photo and GPS uploads the photo to private R2 and writes Oracle report records", "Replace after the Medusa report and photograph flow is tested"),
            ("app/api/report-photos/[photoId]/route.ts", "Checks the signed-in citizen owns a photo then returns its private file as an image", "Keep access checks when changing storage providers"),
            ("lib/auth.ts", "Validates Bearer JWT tokens for protected Next.js APIs", "Change current Next.js authentication rules"),
            ("lib/db.ts", "Creates the Oracle connection pool from environment variables", "Keep until the Oracle rollback path is retired"),
            ("lib/report-storage.ts", "Creates the private R2 client and uploads downloads or deletes photograph objects", "Change R2 credentials bucket or storage implementation"),
            ("next.config.ts", "Allows approved development origins and externalizes the Oracle native package", "Change development hostnames or remove Oracle packaging after cutover"),
            (".env.local", "Stores local Next.js Oracle JWT SMS and upload settings", "Local secret changes only and never commit"),
            ("database/oracle/001_mirror_foundation.sql", "Creates the eight Oracle tables sequences indexes and constraints", "Use only for approved Oracle schema setup or reference"),
            ("scripts/check-oracle.mjs", "Checks Oracle authentication tables columns and sequences", "Run while Oracle remains part of the system"),
            ("scripts/check-r2.mjs", "Checks private R2 credentials and bucket reachability without displaying secrets", "Run after adding or rotating R2 settings"),
            ("backend/apps/backend/medusa-config.ts", "Registers the Mirror module Neon database and Redis modules", "Change Medusa infrastructure configuration"),
            ("backend/apps/backend/src/modules/mirror/models/mirror.ts", "Defines the eight PostgreSQL entities fields relations and indexes", "Change the target data model then generate a migration"),
            ("backend/apps/backend/src/modules/mirror/service.ts", "Provides Medusa create list update and delete operations for Mirror entities", "Extend domain operations when custom logic is required"),
            ("backend/apps/backend/src/modules/mirror/index.ts", "Registers the custom Mirror module", "Change only if the module name or service registration changes"),
            ("backend/apps/backend/src/modules/mirror/migrations", "Contains the generated PostgreSQL migration and snapshot", "Generate from model changes and never rewrite applied history"),
            ("backend/apps/backend/.env", "Stores local Neon Redis CORS and authentication secrets", "Update privately and never commit"),
            ("backend/apps/backend/.env.template", "Documents required variable names without real credentials", "Update whenever a new required variable is introduced"),
            ("package.json and backend/package.json", "Define dependencies and run build test and development commands", "Change scripts versions or dependencies"),
            ("docs/The_Mirror_Project_Developer_Handover.docx", "Provides the current architecture status file map procedures and handover checklist", "Update after every completed infrastructure or application milestone"),
            ("scripts/build_developer_handover.py", "Regenerates this Word document in a consistent format", "Update the documented facts then rerun it and visually verify every page"),
        ],
        [2.25, 3.05, 1.4],
    )

    add_heading(doc, "5 Changes added in the Medusa migration", 1)
    add_table(
        doc,
        ["Path", "Change", "Purpose"],
        [
            ("backend/", "Created MedusaJS 2.20.1 backend workspace", "Keeps the replacement backend separate from the existing Next.js application."),
            ("backend/apps/backend/medusa-config.ts", "Registered the Mirror module and conditional Redis modules", "Loads Neon through DATABASE_URL and Redis through REDIS_URL."),
            ("src/modules/mirror/models/mirror.ts", "Defined eight related data models", "Replaces Oracle number and sequence design with Medusa string IDs and PostgreSQL foreign keys."),
            ("src/modules/mirror/service.ts", "Created the generated Medusa service", "Provides create list update and delete methods for the Mirror models."),
            ("src/modules/mirror/index.ts", "Registered the custom module", "Makes the service available through the Medusa dependency container as mirror."),
            ("src/modules/mirror/migrations/", "Generated and applied the initial migration", "Creates tables constraints indexes and relationships in Neon."),
            ("backend/apps/backend/.env.template", "Documented the rediss connection format", "Shows developers the expected variable without storing a real secret."),
            ("backend/apps/backend/.env", "Contains local secrets and connection strings", "Ignored by Git and must never be copied into source control."),
        ],
        [2.2, 2.15, 2.35],
    )

    add_heading(doc, "Redis modules registered", 2)
    for text in (
        "Redis caching provider",
        "Redis event bus with worker concurrency set to one",
        "Redis workflow engine",
        "Redis distributed locking provider",
    ):
        add_bullet(doc, text)
    doc.add_paragraph(
        "These modules load only when REDIS_URL is present. This allows source checks without Redis, "
        "while production receives persistent coordination instead of Medusa's in-memory fallbacks."
    )

    add_heading(doc, "6 PostgreSQL data model", 1)
    add_table(
        doc,
        ["Table", "Purpose", "Important relationships and fields"],
        [
            ("mirror_user", "Citizen and staff profiles", "Unique phone name required optional email role and phone verification flag"),
            ("mirror_otp_request", "Hashed OTP challenges", "Phone OTP hash expiry attempts verification time and optional user link"),
            ("mirror_session", "Login sessions", "Unique token hash expiry revocation time and required user link"),
            ("mirror_pothole", "Physical road defect record", "Latitude longitude address severity current status and linked reports"),
            ("mirror_report", "Citizen submission", "Public report ID description status submission time citizen and pothole links"),
            ("mirror_report_photo", "Photograph metadata", "Storage key file details capture coordinates confirmation coordinates distance and 100 metre result"),
            ("mirror_status_history", "Status changes over time", "Old status new status note actor identifier time and report link"),
            ("mirror_audit_log", "Security and accountability history", "Actor action entity details IP time and optional report link"),
        ],
        [1.55, 1.9, 3.25],
    )
    doc.add_paragraph(
        "Medusa automatically adds created_at updated_at and deleted_at columns. Records use prefixed "
        "text IDs such as musr and mrpt rather than Oracle NUMBER values and sequences. PostgreSQL "
        "indexes support phone report ID status date and location lookups."
    )

    add_heading(doc, "Photo and location validation design", 2)
    doc.add_paragraph(
        "The target report flow takes the photograph first and captures the current GPS location second. "
        "The backend will calculate the distance between the location recorded near photo capture and the "
        "confirmed current location. The report should be accepted only when the distance is 100 metres or "
        "less and the GPS accuracy and timestamps satisfy the validation rules. The database stores both "
        "coordinate pairs, the calculated distance, and the pass or fail result for later auditing."
    )

    add_heading(doc, "7 Environment variables and secrets", 1)
    add_table(
        doc,
        ["Variable", "Used by", "Required state"],
        [
            ("DATABASE_URL", "Medusa and Neon", "Direct Neon PostgreSQL URL for migrations and a suitable production URL for Railway"),
            ("REDIS_URL", "Medusa and Upstash", "Upstash TCP URL beginning with rediss:// not the REST URL"),
            ("JWT_SECRET", "Authentication", "Long random production secret and the same value wherever token verification occurs"),
            ("COOKIE_SECRET", "Medusa cookies", "Long random production secret different from JWT_SECRET"),
            ("STORE_CORS", "Browser access", "Frontend production domain and approved local development origins"),
            ("ADMIN_CORS", "Medusa administration", "Only approved admin origins"),
            ("AUTH_CORS", "Authentication", "Only approved frontend and admin origins"),
            ("SMS_API_KEY", "Future SMS provider", "Pending while the temporary OTP remains enabled"),
            ("R2_ENDPOINT", "Next.js and Cloudflare R2", "Account S3 endpoint over HTTPS"),
            ("R2_ACCESS_KEY_ID", "Next.js and Cloudflare R2", "Bucket-scoped Object Read and Write key"),
            ("R2_SECRET_ACCESS_KEY", "Next.js and Cloudflare R2", "Secret stored only in private runtime settings"),
            ("R2_BUCKET_NAME", "Next.js and Cloudflare R2", "Private bucket containing report-photos objects"),
        ],
        [1.55, 1.75, 3.4],
    )
    doc.add_paragraph(
        "Local Medusa secrets belong in backend/apps/backend/.env. Root Next.js and Oracle secrets belong "
        "in .env.local. Railway and Vercel secrets belong in each platform's Environment Variables page. "
        "Never place a real URL password token or API key in .env.template documentation screenshots or Git."
    )

    add_heading(doc, "8 Running and verifying the project", 1)
    add_heading(doc, "Run the existing Next.js application", 2)
    add_code(doc, 'cd "C:\\Users\\Uttara ERP\\mirror-project"')
    add_code(doc, "npm install")
    add_code(doc, "npm run db:check")
    add_code(doc, "npm run dev")
    doc.add_paragraph("Open http://localhost:3000. Keep the terminal running while testing.")

    add_heading(doc, "Run the Medusa backend", 2)
    add_code(doc, 'cd "C:\\Users\\Uttara ERP\\mirror-project\\backend"')
    add_code(doc, "npm install")
    add_code(doc, "npm run backend:dev")
    doc.add_paragraph(
        "Medusa normally starts on port 9000. A successful production-style startup should show Redis "
        "connection messages and must not show fake Redis local event bus or in-memory locking warnings."
    )

    add_heading(doc, "Database commands", 2)
    add_code(doc, 'cd "C:\\Users\\Uttara ERP\\mirror-project\\backend\\apps\\backend"')
    add_code(doc, "npm exec medusa db:migrate")
    add_code(doc, "npm exec medusa db:generate mirror")
    doc.add_paragraph(
        "Run db:generate only after changing a model. Review the generated migration before running "
        "db:migrate. Never hand-edit Medusa migration snapshots or committed migration history."
    )

    add_heading(doc, "Photograph storage check", 2)
    add_code(doc, 'cd "C:\\Users\\Uttara ERP\\mirror-project"')
    add_code(doc, "npm run storage:check")
    doc.add_paragraph(
        "Run this after adding the four private R2 variables. A successful check confirms that the "
        "application credentials can reach the private report-photo bucket without printing secrets."
    )

    add_heading(doc, "Checks completed", 2)
    for text in (
        "The cleaned Medusa workspace lint and production build completed successfully on 5 September 2026.",
        "The Medusa build uses a fake Redis instance because build-time credentials are not loaded; runtime Redis must be verified separately.",
        "The Next.js lint and production build completed successfully with the private R2 storage code and authenticated photograph route.",
        "The Neon migration completed successfully.",
        "A read-only information_schema check found all eight mirror tables.",
        "The .env file is ignored by the nested backend Git ignore rules.",
        "The Upstash TLS URL format was verified without exposing its password.",
        "A live Upstash PING returned PONG on 5 September 2026.",
        "The R2 integration compiles, but live bucket access is pending the four private R2 environment variables.",
    ):
        add_bullet(doc, text)

    add_heading(doc, "9 Remaining implementation work", 1)
    add_table(
        doc,
        ["Priority", "Work", "Completion test"],
        [
            ("1", "Implement Medusa OTP and session APIs", "Temporary OTP 123456 supports profile completion and returns a valid session"),
            ("2", "Implement Medusa report APIs and workflows", "Authenticated user can submit and retrieve a report with relational records"),
            ("3", "Implement backend photo and GPS validation", "Invalid age accuracy or distance is rejected and valid metadata is stored"),
            ("4", "Verify and migrate Cloudflare R2 photographs", "New upload and authorized download pass and three legacy objects are migrated"),
            ("5", "Connect the Next.js frontend to Medusa", "Browser no longer calls the Oracle routes for the migrated operations"),
            ("6", "Add automated tests", "Authentication report validation and database failure cases pass"),
            ("7", "Deploy Medusa to Railway", "Railway health check is stable and can reach Neon Upstash and private R2"),
            ("8", "Configure production CORS and frontend URL", "The production domain can call the Railway API without CORS errors"),
            ("9", "Complete end to end mobile testing", "OTP camera HTTPS GPS 100 metre validation submit and dashboard all work"),
            ("10", "Remove Oracle only after acceptance", "Backup and rollback plan approved and production data migration verified"),
        ],
        [0.65, 2.75, 3.55],
    )

    add_heading(doc, "10 Deployment and release process", 1)
    for text in (
        "Complete and test the Medusa APIs locally against Neon and Upstash.",
        "Review git status and confirm that no .env or credential file is staged.",
        "Commit the backend workspace and this handover documentation.",
        "Push the branch to GitHub and confirm the expected files are visible without secrets.",
        "Create a Railway service from the GitHub repository and set the backend root directory.",
        "Add production variables to Railway and deploy Medusa.",
        "Run Railway health and API tests before changing the frontend.",
        "Set the frontend backend URL in Vercel and redeploy the Next.js project.",
        "Test the public domain on desktop and mobile over HTTPS.",
        "Keep Oracle available until the new path has passed acceptance and rollback testing.",
    ):
        add_number(doc, text)

    add_heading(doc, "Safe Git commands", 2)
    add_code(doc, "git status")
    add_code(doc, "git check-ignore -v backend/apps/backend/.env")
    add_code(doc, "git add backend docs/The_Mirror_Project_Developer_Handover.docx scripts/build_developer_handover.py")
    add_code(doc, 'git commit -m "Add Medusa backend foundation and developer handover"')
    add_code(doc, "git push origin main")
    doc.add_paragraph(
        "Before committing, inspect git status and the staged diff. Do not use git add until REDIS_URL and "
        "all other secrets are confirmed ignored. A new developer should use a feature branch for the API "
        "and frontend cutover rather than making unreviewed production changes directly on main."
    )

    add_heading(doc, "11 Developer handover checklist", 1)
    checklist = (
        "Repository access and the active branch are confirmed.",
        "Node.js meets the backend requirement of 20.19 or later or 22.12 or later.",
        "The developer has received secrets through an approved private channel.",
        "Root .env.local and backend .env exist locally and are ignored by Git.",
        "Next.js starts and the existing Oracle db:check result is understood.",
        "Medusa starts and connects to Neon and Upstash.",
        "All eight Mirror PostgreSQL tables are present.",
        "Cloudflare R2 ownership backup retention access-key rotation and deletion policy are documented.",
        "The temporary OTP limitation is understood.",
        "The frontend still uses Oracle until the Medusa API cutover is merged.",
        "Railway and Vercel environment ownership is assigned.",
        "Backup rollback monitoring and incident contacts are agreed before launch.",
    )
    for text in checklist:
        paragraph = doc.add_paragraph(f"[ ] {text}")
        paragraph.paragraph_format.space_after = Pt(3)

    add_heading(doc, "12 Troubleshooting reference", 1)
    add_table(
        doc,
        ["Symptom", "Likely cause", "Action"],
        [
            ("UPSTASH_REDIS FAILED or ECONNREFUSED", "Wrong Redis URL or blocked port", "Use the Upstash rediss TCP URL and retest"),
            ("Fake Redis warning", "REDIS_URL is missing or invalid", "Check backend .env and restart Medusa"),
            ("Neon ECONNRESET", "Temporary network or endpoint issue", "Retry and verify the approved URL and SSL settings"),
            ("Vercel OTP fails", "Production still depends on private Oracle", "Deploy Medusa APIs and switch the frontend endpoint"),
            ("Camera or GPS unavailable", "No HTTPS denied permission or poor accuracy", "Use HTTPS allow permissions and test outdoors"),
            ("Photograph file not found", "R2 object is missing or legacy file was not migrated", "Check the object key bucket and migration inventory"),
            ("CORS error", "Railway does not allow the Vercel origin", "Set STORE_CORS and AUTH_CORS to exact HTTPS origins"),
            ("Backend missing after clone", "The backend folder was never committed", "Verify secrets are ignored then commit and push backend"),
            ("Migration mismatch", "Model and migration history differ", "Generate review and apply one new migration"),
        ],
        [1.8, 2.3, 2.85],
    )

    add_heading(doc, "Ownership and security notes", 2)
    doc.add_paragraph(
        "BigRock supplies DNS only; Vercel hosts the frontend, Railway will host the backend, Neon stores "
        "relational data, Upstash coordinates Medusa, and Cloudflare R2 stores private photos. Use named accounts, multifactor authentication, "
        "and least privilege. Treat photographs and GPS as sensitive: keep storage private, use short-lived "
        "authorized downloads, validate uploads, and define retention and deletion before launch."
    )

    for section in doc.sections:
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_paragraph.add_run("The Mirror Project Complete Development and Handover Guide  |  5 September 2026")
        set_run_font(run, size=8, color="666666")

    doc.core_properties.title = "The Mirror Project Complete Development and Handover Guide"
    doc.core_properties.subject = "System flow services file responsibilities operating procedures and remaining work"
    doc.core_properties.author = "The Mirror Project Team"
    doc.core_properties.keywords = "MedusaJS Neon PostgreSQL Upstash Redis Railway Vercel"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    build_document()
